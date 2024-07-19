import os
import io
import shutil
import zipfile
from PIL import Image

import streamlit as st
from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

# Define functions
def set_font(paragraph, font_name, font_size):
  for run in paragraph.runs:
      run.font.name = font_name
      r = run._element
      rFonts = r.rPr.rFonts
      rFonts.set(qn('w:eastAsia'), font_name)
      run.font.size = Pt(font_size)

def set_document_font(doc, font_name='Times New Roman', font_size=12):
  for paragraph in doc.paragraphs:
      set_font(paragraph, font_name, font_size)
  for table in doc.tables:
      for row in table.rows:
          for cell in row.cells:
              for paragraph in cell.paragraphs:
                  set_font(paragraph, font_name, font_size)

def append_pdf_to_docx(pdf_paths, docx_path):
  try:
      doc = Document(docx_path)
  except:
      doc = Document()

  for pdf_path in pdf_paths:
      pdf_document = fitz.open(pdf_path)

      for page_num in range(len(pdf_document)):
          page = pdf_document.load_page(page_num)
          pix = page.get_pixmap()
          img_data = pix.tobytes()
          image = Image.open(io.BytesIO(img_data))
          image_stream = io.BytesIO()
          image.save(image_stream, format='PNG')
          doc.add_picture(image_stream, width=Inches(6))

          if page_num < len(pdf_document) - 1:
              doc.add_page_break()

  set_document_font(doc)
  doc.save(docx_path)

def append_docx_to_docx(docx_paths, docx_path):
  try:
      doc = Document(docx_path)
  except:
      doc = Document()

  for path in docx_paths:
      sub_doc = Document(path)
      for element in sub_doc.element.body:
          doc.element.body.append(element)

  set_document_font(doc)
  doc.save(docx_path)

def process_zip(zip_path, file_paths, output_zip_path, file_type):
  extract_dir = "extracted_docs"
  os.makedirs(extract_dir, exist_ok=True)

  with zipfile.ZipFile(zip_path, 'r') as zip_ref:
      zip_ref.extractall(extract_dir)

  for root, _, files in os.walk(extract_dir):
      for file in files:
          if file.endswith('.docx'):
              docx_path = os.path.join(root, file)
              if file_type == 'pdf':
                  append_pdf_to_docx(file_paths, docx_path)
              elif file_type == 'docx':
                  append_docx_to_docx(file_paths, docx_path)

  with zipfile.ZipFile(output_zip_path, 'w') as zip_ref:
      for root, _, files in os.walk(extract_dir):
          for file in files:
              full_path = os.path.join(root, file)
              zip_ref.write(full_path, os.path.relpath(full_path, extract_dir))

  shutil.rmtree(extract_dir)

# Streamlit app
st.title("Append PDFs or Word Documents to Word Documents")

# Check if required secrets are present
required_secrets = ['client_id', 'client_secret', 'redirect_url', 'bmac_api_key', 'bmac_link']
missing_secrets = [secret for secret in required_secrets if secret not in st.secrets]

if missing_secrets:
  st.error(f"Missing required secrets: {', '.join(missing_secrets)}. Please check your secrets.toml file or Streamlit Cloud secrets configuration.")
else:
  # Add authentication
  try:
      add_auth()
  except Exception as e:
      st.error(f"Error in authentication: {str(e)}")
      # Debugging: Print the exception
      st.write("Exception:", e)
  else:
      # Check if user is authenticated
      if 'email' in st.session_state:
          st.write(f"Welcome, {st.session_state.email}!")

          # Upload zip file
          zip_file = st.file_uploader("Upload ZIP file containing Word documents", type=["zip"])
          # Choose file type to append
          file_type = st.selectbox("Select file type to append", ["PDF", "Word Document"])
          # Upload multiple files
          file_types = {"PDF": "pdf", "Word Document": "docx"}
          files = st.file_uploader(f"Upload {file_type} files to append", type=[file_types[file_type]], accept_multiple_files=True)

          if st.button("Process"):
              if zip_file and files:
                  zip_path = os.path.join("uploaded_files", zip_file.name)
                  os.makedirs("uploaded_files", exist_ok=True)

                  with open(zip_path, "wb") as f:
                      f.write(zip_file.getbuffer())

                  file_paths = []
                  for file in files:
                      file_path = os.path.join("uploaded_files", file.name)
                      with open(file_path, "wb") as f:
                          f.write(file.getbuffer())
                      file_paths.append(file_path)

                  output_zip_path = os.path.join("uploaded_files", "output.zip")
                  process_zip(zip_path, file_paths, output_zip_path, file_types[file_type])

                  with open(output_zip_path, "rb") as f:
                      st.download_button("Download processed ZIP", f, file_name="processed_documents.zip")

                  st.success(f"{file_type}s appended to Word documents in {output_zip_path} successfully.")
                  shutil.rmtree("uploaded_files")
              else:
                  st.error("Please upload both a ZIP file and at least one file to append.")
      else:
          st.warning("Please log in to access this feature.")
          st.markdown(f"[Subscribe Now]({st.secrets.bmac_link})")

# For debugging purposes, you can uncomment the following line to see the contents of st.secrets
# st.write(st.secrets)
