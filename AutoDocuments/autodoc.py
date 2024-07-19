import streamlit as st
import csv
from docx import Document
import zipfile
from copy import deepcopy
from io import BytesIO

# Function to reset the session state
def reset_state():
    st.session_state.clear()

# Call the reset_state function at the start of the app
reset_state()

# Streamlit app
st.title("Auto Document Generator")

# Information about session history clearing
st.markdown("""
### Important Information
This application is designed with your privacy and data security in mind. Here's how it works:

1. **Session State Reset:** At the start of each session, the application resets its internal session state. This means that any data from previous sessions is cleared, ensuring that no information carries over from one use to the next.

2. **Temporary Data Storage:** When you upload your document template and CSV file, this data is only stored temporarily within the current session. As soon as you close the app or start a new session, this data is automatically deleted.

3. **No Permanent Storage:** The app does not store any of your uploaded files or generated documents permanently. Once you download the generated documents, they are removed from the app's memory.

4. **Security:** This process helps ensure that your sensitive information remains secure and is not accessible after you finish using the app.

In summary, every time you use the app, it starts with a clean slate, clearing all previous session data to protect your privacy and keep your data secure.
""")

# Example CSV content
example_csv = """
grantee_name,grant_number,grantee_street,grantee_citystatezip,award_amount_numerical,convert_numbers_to_words,contact_name,contact_title,contact_number,contact_email,sig_name,sig_title
Example Grantee,12345,123 Example St,Example City, ST 12345,1000,One Thousand,Jane Doe,Director,555-1234,jane.doe@example.com,John Smith,CEO
Another Grantee,67890,456 Another St,Another City, ST 67890,2000,Two Thousand,John Roe,Manager,555-5678,john.roe@example.com,Jane Smith,CFO
"""

# Button to download the example CSV file
st.download_button(
    label="Download Example CSV",
    data=example_csv,
    file_name="example.csv",
    mime="text/csv"
)

# Upload the template document
template_file = st.file_uploader("Upload Document Template (.docx)", type="docx")
if template_file:
    template_document = Document(template_file)

# Upload the CSV file
csv_file = st.file_uploader("Upload Data CSV (.csv)", type="csv")
if csv_file:
    # Open the CSV file in text mode
    try:
        data = list(csv.DictReader(csv_file.getvalue().decode("utf-8").splitlines()))
    except UnicodeDecodeError:
        st.error("There was an error decoding the CSV file. Please ensure it is in UTF-8 format.")

# Input for unique file name prefix
unique_name = st.text_input("Enter a unique name for the generated files:")

# Dropdown for document type selection
document_type = st.selectbox(
    "Select the type of document:",
    ["Award Letter", "Grant Agreement", "Commitment Letter", "Other"]
)

# Function to replace placeholders in the document
def replace_placeholders(document, data):
    for key, value in data.items():
        # Replace placeholders in the paragraphs
        for paragraph in document.paragraphs:
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
        # Replace placeholders in the header
        for section in document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

# Generate documents and create a zip file
if st.button("Generate Documents") and template_file and csv_file and unique_name:
    required_fields = ['grantee_name', 'grant_number', 'grantee_street', 'grantee_citystatezip',
                       'award_amount_numerical', 'convert_numbers_to_words', 'contact_name', 
                       'contact_title', 'contact_number', 'contact_email', 'sig_name', 'sig_title']
    document_paths = []
    zip_buffer = BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w') as zipf:
        for row in data:
            # Check for missing keys before processing
            missing_fields = [field for field in required_fields if field not in row]
            if missing_fields:
                st.error(f"Missing required fields in the CSV file: {', '.join(missing_fields)}")
                continue

            # Create a new document based on the template
            new_document = deepcopy(template_document)

            # Replace placeholders with data from the current row
            replace_placeholders(new_document, row)

            # Define the file path for the new document and change name
            file_name = f"{unique_name}_{row['grantee_name']}_{document_type.replace(' ', '')}.docx"
            doc_buffer = BytesIO()
            new_document.save(doc_buffer)
            doc_buffer.seek(0)

            # Add the document to the zip file
            zipf.writestr(file_name, doc_buffer.read())

    zip_buffer.seek(0)
    st.download_button(
        label="Download Documents",
        data=zip_buffer,
        file_name="documents.zip",
        mime="application/zip"
    )
